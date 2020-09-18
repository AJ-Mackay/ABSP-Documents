# Module 14 - Excel, Word and PDF Documents: Reading and Editing Word Documents

import docx, os
os.chdir('/Users/paulmackay/Desktop/Python/Excel, Word and PDF Documents')

d = docx.Document('demo.docx')

d.paragraphs # Returns a list of paragraph objects.
d.paragraphs[0].text # Returns 'Document Title'
d.paragraphs[1].text # Returns 'A plain paragraph with some bold and some italic'

p = d.paragraphs[1]
p.runs # Returns a list of run objects

### A new Run starts whenever there is a change in the style e.g bold, italic ###
p.runs[0].text # Returns 'A plain paragraph having some '
p.runs[1].text # Returns 'bold'
p.runs[2].text # Returns ' and some '
p.runs[3].text # Returns 'italic.'

p.runs[1].bold # Returns True
p.runs[3].italic # Returns True

p.runs[3].underline = True # This underlines the text.
p.runs[3].text = 'italic and underlined.' # This changes the text.
d.save('demo2.docx')

### SHELL IS PLAYING UP AGAIN ###

### Results ###
### Expected: 'A plain paragraph with some bold and some italic and underlined' ###
### Actual: 'A plain paragraph with some bolditalic and underlined.italic' ###

p.style = 'Title'
d.save('demo3.docx')

d = docx.Document()
d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('This is another paragraph.')
d.save('demo4.docx')

p = d.paragraphs[0]
p.add_run('This is a new run.')
p.runs[1].bold = True
d.save('demo5.docx')

# Full documentation available at python-docx.readthedocs.org
